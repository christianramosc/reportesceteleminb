# -*- coding: utf-8 -*-
"""
Genera la versión .docx (editable en Word) de los reportes.

CÓMO FUNCIONA AHORA (importante si vas a tocar este archivo):
El Word ya NO se redacta aquí. Se construye a partir de la MISMA lista de
elementos que el script del PDF arma para ReportLab: cada script de reporte
acepta `recolectar_elementos=[]` y deja ahí sus flowables; después
flowables_a_docx.volcar_elementos() los traduce a Word en el mismo orden.

Por eso el Word ahora lleva las mismas secciones, los mismos párrafos, las
mismas tablas y las gráficas intercaladas donde van, igual que el PDF.

=> Si quieres cambiar un texto del reporte, cámbialo en el script del PDF
   (avance_preliminar_original.py, comparativo_mensual_original.py o
   resumen_mensual_original.py). El Word hereda el cambio automáticamente.
   NO vuelvas a escribir textos aquí, o volverán a desincronizarse.
"""

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

try:
    from . import flowables_a_docx as _fd
except ImportError:  # ejecución suelta
    import flowables_a_docx as _fd


# =======================================================================
# Documento base
# =======================================================================
def _preparar_documento():
    """Crea el documento en carta (Letter) con los mismos márgenes que el
    PDF, para que las tablas y las gráficas guarden la misma proporción."""
    documento = Document()

    for seccion in documento.sections:
        seccion.page_width = Cm(21.59)
        seccion.page_height = Cm(27.94)
        seccion.left_margin = Cm(2.2)
        seccion.right_margin = Cm(2.2)
        seccion.top_margin = Cm(2.0)
        seccion.bottom_margin = Cm(1.8)

    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(10.5)

    # Los Heading de Word vienen en azul claro por defecto; se alinean con
    # la identidad de los PDFs (azul Inbursa y rojo MG).
    for nombre, color, tamano in (("Heading 1", _fd.INBURSA_AZUL, 15.5),
                                  ("Heading 2", _fd.MG_ROJO_OSCURO, 12.5)):
        try:
            estilo = documento.styles[nombre]
            estilo.font.color.rgb = color
            estilo.font.size = Pt(tamano)
            estilo.font.bold = True
            estilo.font.name = "Calibri"
        except KeyError:
            pass

    return documento


def _agregar_encabezado_pie(documento, nombre_empresa, subtitulo, nombre_analista):
    """Reproduce el encabezado y el pie que el PDF dibuja en cada página."""
    seccion = documento.sections[0]

    enc = seccion.header.paragraphs[0]
    enc.text = ""
    run_empresa = enc.add_run(nombre_empresa)
    run_empresa.bold = True
    run_empresa.font.size = Pt(10.5)
    run_empresa.font.color.rgb = _fd.INBURSA_AZUL
    enc.add_run("\n")
    run_sub = enc.add_run(subtitulo)
    run_sub.font.size = Pt(8)
    run_sub.font.color.rgb = _fd.GRIS_TEXTO_SEC

    pie = seccion.footer.paragraphs[0]
    pie.text = ""
    run_pie = pie.add_run(
        f"{nombre_analista}    ·    {nombre_empresa}    ·    "
        f"Generado el {date.today().strftime('%d/%m/%Y')}"
    )
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = _fd.GRIS_TEXTO_SEC
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _construir(elementos, nombre_archivo, carpeta_salida, subtitulo,
               nombre_pdv, nombre_analista, nombre_empresa="AUTOEXPRESS INBURSA"):
    """Tronco común de los tres reportes: documento base + encabezado/pie +
    volcado de los flowables del PDF + nota final."""
    documento = _preparar_documento()
    _agregar_encabezado_pie(
        documento, nombre_empresa, f"{subtitulo} · {nombre_pdv}", nombre_analista
    )
    _fd.volcar_elementos(documento, elementos)

    os.makedirs(carpeta_salida, exist_ok=True)
    ruta_final = os.path.join(carpeta_salida, nombre_archivo)
    documento.save(ruta_final)
    return ruta_final


# =======================================================================
# Reporte: Avance Preliminar Mensual
# =======================================================================
def generar_docx_avance(elementos, nombre_pdv, nombre_analista,
                        nombre_archivo=None, carpeta_salida="."):
    if nombre_archivo is None:
        nombre_archivo = f"Avance_Preliminar_{date.today().strftime('%Y%m%d_%H%M')}.docx"
    return _construir(elementos, nombre_archivo, carpeta_salida,
                      "Avance Preliminar de Bitácora", nombre_pdv, nombre_analista)


# =======================================================================
# Reporte: Comparativo Mensual
# =======================================================================
def generar_docx_comparativo(elementos, nombre_pdv, nombre_analista,
                             nombre_archivo=None, carpeta_salida="."):
    if nombre_archivo is None:
        nombre_archivo = f"Reporte_Comparativo_{date.today().strftime('%Y%m%d_%H%M')}.docx"
    return _construir(elementos, nombre_archivo, carpeta_salida,
                      "Reporte Comparativo Mensual", nombre_pdv, nombre_analista)


# =======================================================================
# Reporte: Resumen Mensual (antes no tenía versión Word)
# =======================================================================
def generar_docx_resumen(elementos, nombre_pdv, nombre_analista,
                         nombre_archivo=None, carpeta_salida="."):
    if nombre_archivo is None:
        nombre_archivo = f"Resumen_Mensual_{date.today().strftime('%Y%m%d_%H%M')}.docx"
    return _construir(elementos, nombre_archivo, carpeta_salida,
                      "Reporte Mensual de Bitácora", nombre_pdv, nombre_analista)
