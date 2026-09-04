# -*- coding: utf-8 -*-
"""
TRADUCTOR PDF -> WORD

===========================================================================
 ¿POR QUÉ EXISTE ESTE ARCHIVO?
===========================================================================
Antes, el .docx se generaba con código propio, aparte del PDF: tenía sus
propios títulos, sus propios textos y volcaba todas las gráficas juntas al
final. Resultado: el Word no se parecía al PDF, y cada vez que se editaba
un texto del PDF había que acordarse de editarlo también en el Word (cosa
que nunca pasa).

Aquí se hace lo contrario: el Word se construye a partir de la MISMA lista
de elementos (`elementos`) que los scripts arman para ReportLab. Cada
Paragraph, Table e Image del PDF se traduce a su equivalente de Word, en el
mismo orden.

Consecuencia práctica: **el texto del reporte se escribe UNA sola vez, en
el script del PDF.** Si mañana cambias un párrafo ahí, el Word sale con ese
cambio solo. No hay nada que sincronizar a mano.

El Word no queda idéntico píxel a píxel (Word no es un motor de maquetación
como ReportLab), pero sí lleva las mismas secciones, en el mismo orden, con
los mismos textos, tablas y gráficas — y con estilos nativos de Word para
que se pueda editar con comodidad.
"""

import os
import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Paleta (misma de los PDFs)
MG_ROJO         = RGBColor(0xE4, 0x00, 0x2B)
MG_ROJO_OSCURO  = RGBColor(0x7A, 0x00, 0x19)
INBURSA_AZUL    = RGBColor(0x19, 0x19, 0x70)
GRIS_OSCURO     = RGBColor(0x2B, 0x2B, 0x2B)
GRIS_TEXTO_SEC  = RGBColor(0x6B, 0x72, 0x80)

HEX_INBURSA_AZUL = "191970"
HEX_AZUL_PALIDO  = "EEF0F8"
HEX_GRIS_ZEBRA   = "F6F7F9"

# Ancho útil de la página carta con los márgenes que usa _preparar_documento
ANCHO_UTIL_CM = 17.0


# ===========================================================================
#  1) Texto: del marcado de ReportLab a runs de Word
# ===========================================================================
# ReportLab acepta un mini-HTML en sus Paragraph: <b>, <i>, <font color=...>,
# <br/>, &nbsp;. Word no entiende nada de eso, así que hay que convertirlo a
# "runs" (fragmentos con formato) en vez de imprimir las etiquetas tal cual.
_TOKEN = re.compile(
    r"(<b>|</b>|<i>|</i>|<br\s*/?>|<font[^>]*>|</font>|<super>|</super>)",
    re.IGNORECASE,
)
_COLOR_EN_FONT = re.compile(r"color\s*=\s*['\"]?#?([0-9A-Fa-f]{6})", re.IGNORECASE)


def _desescapar(texto):
    return (texto.replace("&nbsp;", " ")
                 .replace("&amp;", "&")
                 .replace("&lt;", "<")
                 .replace("&gt;", ">"))


def escribir_texto_con_formato(parrafo, texto, tamano=None, color=None,
                               negrita_base=False, italica_base=False):
    """Vuelca `texto` (con marcado de ReportLab) dentro de un párrafo de Word,
    respetando negritas, itálicas, saltos de línea y colores."""
    if texto is None:
        return parrafo

    negrita = negrita_base
    italica = italica_base
    color_actual = color
    pila_color = []

    for parte in _TOKEN.split(str(texto)):
        if not parte:
            continue
        etiqueta = parte.lower()
        if etiqueta == "<b>":
            negrita = True
        elif etiqueta == "</b>":
            negrita = negrita_base
        elif etiqueta in ("<i>", "<super>"):
            italica = True
        elif etiqueta in ("</i>", "</super>"):
            italica = italica_base
        elif etiqueta.startswith("<br"):
            parrafo.add_run().add_break()
        elif etiqueta.startswith("<font"):
            pila_color.append(color_actual)
            encontrado = _COLOR_EN_FONT.search(parte)
            if encontrado:
                color_actual = RGBColor.from_string(encontrado.group(1).upper())
        elif etiqueta == "</font>":
            color_actual = pila_color.pop() if pila_color else color
        else:
            run = parrafo.add_run(_desescapar(parte))
            run.bold = negrita
            run.italic = italica
            if tamano is not None:
                run.font.size = Pt(tamano)
            if color_actual is not None:
                run.font.color.rgb = color_actual
    return parrafo


def texto_plano(valor):
    """Extrae el texto de un Paragraph de ReportLab (o de un str), ya sin
    etiquetas — se usa para las celdas de las tablas."""
    if valor is None:
        return ""
    crudo = getattr(valor, "text", valor)
    sin_tags = re.sub(r"<[^>]+>", " ", str(crudo))
    return re.sub(r"\s+", " ", _desescapar(sin_tags)).strip()


# ===========================================================================
#  2) Utilidades de formato de Word
# ===========================================================================
def _sombrear(celda, hex_color):
    tc_pr = celda._tc.get_or_add_tcPr()
    sombra = tc_pr.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color}
    )
    tc_pr.append(sombra)


def _sin_bordes(tabla):
    """Quita los bordes de una tabla que solo sirve de maquetación (las
    tarjetas KPI y el recuadro de introducción)."""
    tbl_pr = tabla._tbl.tblPr
    bordes = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elemento = bordes.makeelement(qn(f"w:{lado}"), {qn("w:val"): "none"})
        bordes.append(elemento)
    tbl_pr.append(bordes)


def _es_flowable(objeto, nombre_clase):
    return type(objeto).__name__ == nombre_clase


def _nombre_estilo(flowable):
    estilo = getattr(flowable, "style", None)
    return getattr(estilo, "name", "") or ""


def _formato_del_estilo(flowable):
    """Lee tamaño, color, negritas e itálicas DIRECTAMENTE del estilo de
    ReportLab. Así el Word hereda el formato real del PDF y no depende de
    que alguien mantenga una tabla de equivalencias a mano (que es lo que
    fallaba con los estilos clonados, como el de las tarjetas KPI)."""
    estilo = getattr(flowable, "style", None)
    if estilo is None:
        return {}
    fuente = (getattr(estilo, "fontName", "") or "").lower()
    color_rl = getattr(estilo, "textColor", None)
    hex_color = None
    if color_rl is not None:
        try:
            hex_color = color_rl.hexval()[2:].upper()   # '0xRRGGBB' -> 'RRGGBB'
        except Exception:
            hex_color = None
    return {
        "tamano": getattr(estilo, "fontSize", None),
        "color": hex_color,
        "negrita": "bold" in fuente,
        "italica": ("italic" in fuente) or ("oblique" in fuente),
        "alineacion": getattr(estilo, "alignment", None),
    }


# ===========================================================================
#  3) Traducción de cada tipo de flowable
# ===========================================================================
# El tamaño, el color y las negritas ya vienen capturados del propio estilo
# de ReportLab (ver _formato_del_estilo). Aquí solo queda decidir qué estilos
# se mapean a los Heading nativos de Word (para que salgan en el panel de
# navegación) y cuáles van justificados.
_HEADINGS = {
    "H1MG": "Heading 1",
    "H2MG": "Heading 2",
    "TituloPortada": "Title",
}
_JUSTIFICADOS = {"CuerpoMG"}
_CENTRADOS = {"KPIValor", "KPILabel"}


def _color_word(hex_color, por_defecto=GRIS_OSCURO):
    if not hex_color:
        return por_defecto
    try:
        return RGBColor.from_string(hex_color)
    except Exception:
        return por_defecto


def _agregar_parrafo(documento, bloque):
    nombre = bloque.get("estilo", "")

    estilo_word = _HEADINGS.get(nombre)
    parrafo = (documento.add_paragraph(style=estilo_word) if estilo_word
               else documento.add_paragraph())

    escribir_texto_con_formato(
        parrafo, bloque.get("texto", ""),
        tamano=bloque.get("tamano"),
        color=_color_word(bloque.get("color")),
        negrita_base=bloque.get("negrita", False),
        italica_base=bloque.get("italica", False),
    )

    if nombre in _CENTRADOS:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif nombre in _JUSTIFICADOS:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    formato = parrafo.paragraph_format
    formato.space_after = Pt(6)
    if estilo_word:
        formato.space_before = Pt(12)
    return parrafo


def _es_fila_encabezado(fila):
    return any(c.get("estilo", "").startswith("EncabezadoTabla") for c in fila)


def _agregar_tabla(documento, bloque):
    filas = bloque.get("filas") or []
    if not filas:
        return None
    anchos = bloque.get("anchos")

    # --- Caso A: recuadro de una sola celda (el callout azul de la portada) ---
    if len(filas) == 1 and len(filas[0]) == 1 and filas[0][0].get("tipo") != "tabla":
        tabla = documento.add_table(rows=1, cols=1)
        celda = tabla.rows[0].cells[0]
        celda.paragraphs[0].text = ""
        escribir_texto_con_formato(
            celda.paragraphs[0], filas[0][0].get("texto", ""),
            tamano=10, color=GRIS_OSCURO,
        )
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sombrear(celda, HEX_AZUL_PALIDO)
        documento.add_paragraph()
        return tabla

    # --- Caso B: fila de tarjetas KPI (celdas que son, a su vez, tablas) ---
    if any(c.get("tipo") == "tabla" for fila in filas for c in fila):
        tarjetas = [c for c in filas[0] if c.get("tipo") == "tabla"]
        if not tarjetas:
            return None
        tabla = documento.add_table(rows=1, cols=len(tarjetas))
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        for celda_word, tarjeta in zip(tabla.rows[0].cells, tarjetas):
            celda_word.paragraphs[0].text = ""
            primero = True
            for fila_interna in tarjeta.get("filas", []):
                for valor in fila_interna:
                    texto = valor.get("texto", "")
                    if not texto_plano(texto):
                        continue
                    parrafo = (celda_word.paragraphs[0] if primero
                               else celda_word.add_paragraph())
                    primero = False
                    escribir_texto_con_formato(
                        parrafo, texto,
                        tamano=valor.get("tamano"),
                        color=_color_word(valor.get("color")),
                        negrita_base=valor.get("negrita", False),
                    )
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    parrafo.paragraph_format.space_after = Pt(0)
            _sombrear(celda_word, HEX_GRIS_ZEBRA)
        documento.add_paragraph()
        return tabla

    # --- Caso C: tabla de datos normal ---
    n_cols = max(len(f) for f in filas)
    tabla = documento.add_table(rows=0, cols=n_cols)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    hay_encabezado = _es_fila_encabezado(filas[0])
    for i, fila in enumerate(filas):
        celdas = tabla.add_row().cells
        es_encabezado = (i == 0 and hay_encabezado)
        for j in range(n_cols):
            valor = fila[j] if j < len(fila) else {"texto": ""}
            celda = celdas[j]
            celda.text = ""
            parrafo = celda.paragraphs[0]
            run = parrafo.add_run(texto_plano(valor.get("texto", "")))
            run.font.size = Pt(8.5 if es_encabezado else 8)
            if es_encabezado:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _sombrear(celda, HEX_INBURSA_AZUL)
            elif i % 2 == 0:
                _sombrear(celda, HEX_GRIS_ZEBRA)
            parrafo.alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                                 else WD_ALIGN_PARAGRAPH.CENTER)
            parrafo.paragraph_format.space_after = Pt(0)

    # Respeta las proporciones de ancho que trae la tabla del PDF
    if anchos and len(anchos) == n_cols and all(anchos):
        total = float(sum(anchos))
        for j, ancho in enumerate(anchos):
            ancho_cm = Cm(ANCHO_UTIL_CM * float(ancho) / total)
            for fila_word in tabla.rows:
                fila_word.cells[j].width = ancho_cm

    documento.add_paragraph()
    return tabla


def _agregar_imagen(documento, bloque):
    ruta = bloque.get("ruta")
    if not isinstance(ruta, str) or not os.path.exists(ruta):
        return None
    # Convierte el ancho en puntos que usa el PDF a centímetros de Word,
    # con tope para que nunca se salga del margen.
    ancho_pt = bloque.get("ancho_pt")
    ancho_cm = min(ANCHO_UTIL_CM, (ancho_pt / 28.3465) if ancho_pt else 15.5)
    documento.add_picture(ruta, width=Cm(ancho_cm))
    documento.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return documento.paragraphs[-1]


# ===========================================================================
#  4) Captura del "guion" ANTES de maquetar el PDF
# ===========================================================================
# OJO — esto es sutil y es la razón de que exista esta sección:
# ReportLab MUTA los flowables cuando maqueta. En particular, al hacer
# doc.build() las celdas de una Table dejan de ser Paragraph y se vuelven
# objetos internos (_ExpandedCellTuple), perdiendo el atributo .text.
# Por eso NO se puede guardar la lista de flowables y leerla después: hay
# que extraer el contenido a estructuras de datos simples ANTES del build.
# Eso es lo que hace capturar_guion(); volcar_elementos() ya trabaja sobre
# esos datos, no sobre objetos de ReportLab.
def _desenvolver(valor):
    """Las celdas ya maquetadas vienen envueltas en tuplas/listas."""
    while isinstance(valor, (list, tuple)) and len(valor) == 1:
        valor = valor[0]
    return valor


def _capturar_celda(valor):
    valor = _desenvolver(valor)
    tipo = type(valor).__name__
    if tipo == "Table":
        return _capturar_tabla(valor)
    if tipo == "Paragraph":
        bloque = {"tipo": "parrafo", "estilo": _nombre_estilo(valor),
                  "texto": getattr(valor, "text", "")}
        bloque.update(_formato_del_estilo(valor))
        return bloque
    if tipo == "Image":
        return {"tipo": "imagen", "ruta": getattr(valor, "filename", None),
                "ancho_pt": getattr(valor, "drawWidth", None)}
    return {"tipo": "texto", "texto": "" if valor is None else str(valor)}


def _capturar_tabla(tabla_rl):
    filas = getattr(tabla_rl, "_cellvalues", None) or []
    anchos = getattr(tabla_rl, "_colWidths", None) or getattr(tabla_rl, "_argW", None)
    return {
        "tipo": "tabla",
        "filas": [[_capturar_celda(c) for c in fila] for fila in filas],
        "anchos": [float(a) if a else 0.0 for a in anchos] if anchos else None,
    }


def capturar_guion(elementos):
    """Convierte la lista de flowables de ReportLab en datos simples
    (listas y diccionarios) que sobreviven al doc.build()."""
    guion = []
    for flowable in elementos:
        tipo = type(flowable).__name__
        if tipo == "Paragraph":
            bloque = {"tipo": "parrafo", "estilo": _nombre_estilo(flowable),
                      "texto": getattr(flowable, "text", "")}
            bloque.update(_formato_del_estilo(flowable))
            guion.append(bloque)
        elif tipo == "Table":
            guion.append(_capturar_tabla(flowable))
        elif tipo == "Image":
            guion.append({"tipo": "imagen", "ruta": getattr(flowable, "filename", None),
                          "ancho_pt": getattr(flowable, "drawWidth", None)})
        elif tipo == "PageBreak":
            guion.append({"tipo": "salto"})
        elif tipo == "HRFlowable":
            guion.append({"tipo": "regla"})
        elif tipo == "KeepTogether":
            guion.extend(capturar_guion(getattr(flowable, "_content", []) or []))
        # Spacer y demás flowables de maquetación se omiten: Word maneja su
        # propio espaciado entre párrafos.
    return guion


# ===========================================================================
#  5) Volcado del guion al documento de Word
# ===========================================================================
def volcar_elementos(documento, guion):
    """Escribe en Word el guion capturado del PDF, en el mismo orden."""
    for bloque in guion:
        tipo = bloque.get("tipo")

        if tipo == "parrafo":
            _agregar_parrafo(documento, bloque)
        elif tipo == "tabla":
            _agregar_tabla(documento, bloque)
        elif tipo == "imagen":
            _agregar_imagen(documento, bloque)
        elif tipo == "salto":
            documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif tipo == "regla":
            # La regla horizontal del PDF no hace falta en Word: los
            # encabezados ya se distinguen por su estilo propio.
            pass
        elif tipo == "texto" and bloque.get("texto", "").strip():
            documento.add_paragraph(bloque["texto"])
    return documento
